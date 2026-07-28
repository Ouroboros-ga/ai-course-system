"""PythonDocx parser provider.

Wraps ``python-docx`` to extract paragraphs, headings, and tables, producing
a ``ParserOutput`` compatible with the DocumentIR mapper.

This provider extracts the semantic structure (paragraphs, headings, tables)
but does NOT produce page coordinates on its own -- DOCX has no inherent
pagination. For page coordinates + OCR of embedded images, the pipeline
converts the DOCX to PDF via ``LibreOfficeHeadlessConverter`` and runs the
PDF + PaddleOCR chain (see planner._plan_docx).

If python-docx is not installed, ``parse`` raises ``ParseUnavailableError``;
it never fabricates output.

See docs/phase1/decisions/2026-07-27-统一课程建设九步实施计划.md §7 Step 3.
"""
from __future__ import annotations

import io
import time
import zipfile
from typing import Any, Dict, List, Tuple
from xml.etree import ElementTree

from ..registry import (
    ParserCapabilities,
    ParserOutput,
    ParseUnavailableError,
)
from ..planner import ParsePlan
from ..source_artifact import SourceArtifact

try:
    import docx as _docx  # python-docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    _docx = None  # type: ignore


_WORD_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def _extract_docx_ooxml_supplements(data: bytes) -> tuple[list[dict[str, str]], list[str]]:
    """Read floating-image metadata without inventing page geometry.

    Floating drawings and OMML are outside python-docx's normal body iterator.
    This extraction leaves each item as an exact OOXML locator and asks for
    review where formulas cannot be faithfully normalized to LaTeX.
    """
    supplements: list[dict[str, str]] = []
    warnings: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            if "word/document.xml" not in archive.namelist():
                return supplements, ["DOCX_OOXML_DOCUMENT_MISSING"]
            root = ElementTree.fromstring(archive.read("word/document.xml"))
            for index, drawing in enumerate(root.findall(".//w:drawing", _WORD_NS), start=1):
                doc_pr = drawing.find(".//wp:docPr", _WORD_NS)
                description = str(doc_pr.get("descr", "")).strip() if doc_pr is not None else ""
                title = str(doc_pr.get("title", "")).strip() if doc_pr is not None else ""
                if description or title:
                    supplements.append({
                        "kind": "image_alt_text",
                        "text": description or title,
                        "raw_locator": f"word/document.xml#/w:body//w:drawing[{index}]/wp:docPr",
                    })
                if drawing.find(".//wp:anchor", _WORD_NS) is not None:
                    warnings.append(f"DOCX_FLOATING_DRAWING_REVIEW_REQUIRED:word/document.xml#/w:body//w:drawing[{index}]")
            formula_count = len(root.findall(".//m:oMath", _WORD_NS))
            if formula_count:
                warnings.append(f"DOCX_OMML_FORMULA_REVIEW_REQUIRED:word/document.xml#{formula_count}")
    except (OSError, ValueError, zipfile.BadZipFile, ElementTree.ParseError) as exc:
        warnings.append(f"DOCX_OOXML_SUPPLEMENT_UNAVAILABLE:{type(exc).__name__}")
    return supplements, warnings


class PythonDocxProvider:
    """Parser provider for DOCX files using python-docx."""

    name = "python-docx"
    version = "1.0.0"
    capabilities = ParserCapabilities(
        supported_formats=("docx",),
        supports_tables=True,
        supports_ocr=False,
        supports_notes=False,
        supports_reading_order=True,
        supports_heading_detection=True,
        supports_visual_assets=False,
        supports_coordinates=False,
        requires_gpu=False,
        max_file_size_bytes=50 * 1024 * 1024,
        max_pages=500,
    )

    def __init__(self) -> None:
        self._available = HAS_DOCX

    @property
    def is_real(self) -> bool:
        return True

    async def parse(self, source: SourceArtifact, plan: ParsePlan) -> ParserOutput:
        if not self._available:
            raise ParseUnavailableError(
                "python-docx is not available. Install with: pip install python-docx"
            )
        start = time.perf_counter()

        data = self._get_source_data(source)
        if data is None:
            raise ParseUnavailableError(
                "Source artifact has no accessible data for DOCX parsing"
            )

        try:
            document = _docx.Document(io.BytesIO(data))
        except Exception as exc:
            raise ParseUnavailableError(
                f"python-docx failed to open document: {exc}"
            ) from exc
        supplements, supplement_warnings = _extract_docx_ooxml_supplements(data)

        # DOCX has no reliable native page geometry.  Treat body content as a
        # section unit, retaining OOXML locators rather than inventing page 1.
        text_blocks: List[Dict[str, Any]] = []
        order = 0
        table_count = 0
        # python-docx 1.1+ retains the body order. Fall back defensively for
        # older installations, preserving the explicit limitation in metadata.
        inner = getattr(document, "iter_inner_content", None)
        items = list(inner()) if callable(inner) else list(document.paragraphs) + list(document.tables)
        for body_index, item in enumerate(items, start=1):
            if hasattr(item, "paragraph_format"):
                text = (item.text or "").strip()
                if not text:
                    continue
                style_name = (item.style.name or "").lower() if item.style else ""
                is_heading = "heading" in style_name or "title" in style_name
                text_blocks.append({
                    "text": text,
                    "bbox": None,
                    "confidence": 1.0,
                    "kind": "heading" if is_heading else "paragraph",
                    "heading_level": next((int(ch) for ch in style_name if ch.isdigit()), 1) if is_heading else None,
                    "style_hints": {"style_name": style_name, "native_locator": f"word/document.xml#/w:body/w:p[{body_index}]"},
                    "raw_locator": f"word/document.xml#/w:body/w:p[{body_index}]",
                    "order_index": order,
                })
                order += 1
            else:
                table_count += 1
                for row_index, row in enumerate(item.rows, start=1):
                    for cell_index, cell in enumerate(row.cells, start=1):
                        text = (cell.text or "").strip()
                        if text:
                            text_blocks.append({
                                "text": text,
                                "bbox": None,
                                "confidence": 1.0,
                                "kind": "table",
                                "style_hints": {"native_locator": f"word/document.xml#/w:body/w:tbl[{body_index}]/w:tr[{row_index}]/w:tc[{cell_index}]"},
                                "raw_locator": f"word/document.xml#/w:body/w:tbl[{body_index}]/w:tr[{row_index}]/w:tc[{cell_index}]",
                                "order_index": order,
                            })
                            order += 1

        pages: List[Dict[str, Any]] = [{
            "section_index": 1,
            "text_blocks": text_blocks,
            "tables": [],
            "formulas": [],
            "ooxml_supplements": supplements,
            "coordinate_unit": "native-structure",
        }]

        elapsed = (time.perf_counter() - start) * 1000
        return ParserOutput(
            provider=self.name,
            provider_version=self.version,
            pages=tuple(pages),
            metadata={
                "section_count": 1,
                "duration_ms": int(elapsed),
                "is_fake": False,
                "docx_engine": "python-docx",
                "paragraph_count": order,
                "table_count": table_count,
            },
            warnings=supplement_warnings,
        )

    @staticmethod
    def _get_source_data(source: SourceArtifact) -> Any:
        if source.data is not None:
            return source.data
        if not source.uri:
            return None
        try:
            from app.services.object_storage import get_object_storage
            return get_object_storage().get(source.uri)
        except FileNotFoundError:
            return None
        except Exception:
            return None


def map_docx_output_to_ir(
    output: ParserOutput,
    source: SourceArtifact,
    run_id: str,
    parser_run_id: str,
    normalization_version: str = "1",
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
    """Map ordered DOCX body content without inventing page coordinates."""
    from ..document_ir.models import ContentBlock, Provenance

    blocks: List[Dict[str, Any]] = []
    for section in output.pages:
        section_index = int(section.get("section_index", 1))
        for index, raw in enumerate(section.get("text_blocks", [])):
            text = raw.get("text", "")
            locator = raw.get("raw_locator") or f"word/document.xml#/body/item[{index + 1}]"
            blocks.append({
                "block_id": f"blk_docx_s{section_index}_i{index}",
                "block_type": raw.get("kind", "paragraph"),
                "text": text,
                "page_or_slide": None,
                "bbox": None,
                "reading_order": raw.get("order_index", index),
                "heading_level": raw.get("heading_level"),
                "style_hints": raw.get("style_hints") or {},
                "provenance": Provenance(
                    artifact_id=source.artifact_id, run_id=run_id, parser_run_id=parser_run_id,
                    provider=output.provider, raw_locator=locator, confidence=raw.get("confidence"),
                ),
            })
        for supplement_index, supplement in enumerate(section.get("ooxml_supplements", [])):
            text = supplement.get("text", "").strip()
            if not text:
                continue
            locator = supplement["raw_locator"]
            blocks.append(ContentBlock(
                block_id=f"blk_docx_s{section_index}_xml{supplement_index}",
                block_type="image",
                text=text,
                reading_order=len(blocks),
                style_hints={"native_locator": locator, "bbox_unavailable": True},
                provider="python-docx-ooxml",
                provenance=(Provenance(
                    artifact_id=source.artifact_id, run_id=run_id, parser_run_id=parser_run_id,
                    provider="python-docx-ooxml", raw_locator=locator, confidence=1.0,
                ),),
            ).to_dict())
    return blocks, [], []
